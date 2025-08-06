# backend/app/routes/layout_routes.py

import os
import re
import json
from flask import Blueprint, request, jsonify, current_app
from app.utils.decorators import token_required 
from werkzeug.utils import secure_filename
from app.extensions import db
from app.models import Layout
import docx
import google.generativeai as genai

layout_bp = Blueprint('layout_bp', __name__)

def allowed_file(filename):
    # We continue to allow only docx as it is the most structured
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'docx'}

# --- STEP 1: RAW PARSER (Unchanged) ---
def parse_docx_raw(file_path):
    """Parses a docx file into a raw data structure without interpretation."""
    doc = docx.Document(file_path)
    raw_structure = {'paragraphs': [], 'tables': []}
    for p in doc.paragraphs:
        if p.text.strip():
            raw_structure['paragraphs'].append(p.text.strip())
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        raw_structure['tables'].append(table_data)
    return raw_structure

# --- STEP 2: ADVANCED AI-BASED AGENT PARSER (Heavily Upgraded) ---
def parser_agent_analyze_layout(raw_structure, document_type, jenjang=None, mapel=None):
    """
    Uses Gemini to perform a deep, component-based analysis of an educational document.
    """
    google_api_key = os.getenv('GOOGLE_API_KEY')
    if not google_api_key:
        raise ValueError("GOOGLE_API_KEY not found in .env")
    genai.configure(api_key=google_api_key)

    prompt = f"""
Anda adalah AI ahli dalam dekonstruksi dan analisis dokumen kurikulum pendidikan di Indonesia.
Tugas Anda adalah menganalisis struktur JSON mentah dari sebuah dokumen (`{document_type}`) 
dan mengubahnya menjadi sebuah template komponen JSON yang terstruktur dan cerdas.

# INPUT (Struktur Dokumen Mentah):
{json.dumps(raw_structure, indent=2)}

# KONTEKS:
- Jenjang Pendidikan: {jenjang}
- Mata Pelajaran: {mapel}
- Tipe Dokumen untuk dianalisis: {document_type}
- Dokumen ini bisa berupa Modul Ajar, ATP (Alur Tujuan Pembelajaran), Prota (Program Tahunan), atau Promes (Program Semester).

# PERINTAH UTAMA:
1. Analisis struktur input secara mendalam.
2. Hasilkan sebuah objek JSON tunggal dan valid sebagai output akhir.
3. **Tambahkan metadata** berikut ke JSON agar dapat disimpan di database:
   {{
     "jenjang": "{jenjang}",
     "mapel": "{mapel}",
     "tipe_dokumen": "{document_type}"
   }}

# STRUKTUR JSON YANG DIHARAPKAN:
{{
  "metadata": {{
    "jenjang": "{jenjang}",
    "mapel": "{mapel}",
    "tipe_dokumen": "{document_type}",
    "info_lain": {{ ... jika ada ... }}
  }},
  "document_structure": [
    {{
      "type": "paragraph",
      "text_preview": "contoh kalimat awal...",
      "style": {{
        "alignment": "justify",
        "font_size": 12,
        "font_style": ["bold"],
        "line_spacing": 1.5,
        "indentation": "first-line"
      }}
    }},
    {{
      "type": "table",
      "rows": 4,
      "columns": 3,
      "header_style": {{
        "bold": true,
        "alignment": "center",
        "background_color": "light_gray"
      }},
      "body_style": {{
        "alignment": "left",
        "font_size": 11
      }}
    }}
  ],
  "headers": [
    {{
      "text": "Judul Dokumen",
      "style": {{ "alignment": "center", "font_size": 16, "bold": true }}
    }}
  ]
}}
"""

    try:
        generation_config = genai.GenerationConfig(response_mime_type="application/json")
        model = genai.GenerativeModel('gemini-2.5-flash', generation_config=generation_config)
        response = model.generate_content(prompt)

        return json.loads(response.text, strict=False)

    except Exception as e:
        print(f"[PARSER_AGENT_ERROR] Failed to analyze layout with AI: {e}")
        raise

# --- UPLOAD ROUTE (Unchanged from previous update) ---
@layout_bp.route('/upload', methods=['POST'])
@token_required
def upload_layout(current_user):
    if 'file' not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files['file']
    jenjang = request.form.get('jenjang')
    mapel = request.form.get('mapel')
    tipe_dokumen = request.form.get('tipe_dokumen')
    
    if not all([file, file.filename, jenjang, mapel, tipe_dokumen]):
        return jsonify({"msg": "Missing required form data"}), 400

    if allowed_file(file.filename):
        filename = secure_filename(file.filename)
        upload_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'layouts')
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)

        try:
            raw_layout = parse_docx_raw(file_path)
            smart_template_json = parser_agent_analyze_layout(raw_layout, tipe_dokumen)
        except Exception as e:
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({"msg": f"Error: AI failed to process the layout: {str(e)}"}), 500

        new_layout = Layout(
            jenjang=jenjang, mapel=mapel, tipe_dokumen=tipe_dokumen,
            layout_json=smart_template_json,
            uploaded_by=current_user.id, file_path=file_path
        )
        db.session.add(new_layout)
        db.session.commit()

        return jsonify({
            "msg": "Layout uploaded and successfully analyzed by AI into a smart template.",
            "layout_id": new_layout.id,
            "smart_template": smart_template_json
        }), 201

    return jsonify({"msg": "File type not allowed. Only .docx is supported."}), 400

@layout_bp.route('/', methods=['GET'])
@token_required
def get_all_layouts(current_user):
    """Mengambil daftar semua layout yang diunggah oleh pengguna."""
    layouts = Layout.query.filter_by(uploaded_by=current_user.id).order_by(Layout.created_at.desc()).all()
    
    # Membuat daftar hasil untuk ditampilkan
    result = []
    for layout in layouts:
        result.append({
            'id': layout.id,
            'jenjang': layout.jenjang,
            'mapel': layout.mapel,
            'tipe_dokumen': layout.tipe_dokumen,
            'file_name': os.path.basename(layout.file_path),
            'created_at': layout.created_at.isoformat()
        })
        
    return jsonify(result), 200

# (READ) - Mengambil satu layout spesifik berdasarkan ID
@layout_bp.route('/<int:layout_id>', methods=['GET'])
@token_required
def get_layout_by_id(current_user, layout_id):
    """Mengambil detail satu layout spesifik."""
    layout = Layout.query.get_or_404(layout_id)
    
    # Memastikan pengguna hanya bisa mengakses layout miliknya
    if layout.uploaded_by != current_user.id:
        return jsonify({"msg": "Akses ditolak. Anda bukan pemilik layout ini."}), 403
        
    return jsonify({
        'id': layout.id,
        'jenjang': layout.jenjang,
        'mapel': layout.mapel,
        'tipe_dokumen': layout.tipe_dokumen,
        'layout_json': layout.layout_json, # Mengirimkan JSON hasil analisis AI
        'file_name': os.path.basename(layout.file_path),
        'created_at': layout.created_at.isoformat()
    }), 200

# (UPDATE) - Memperbarui metadata sebuah layout
@layout_bp.route('/<int:layout_id>', methods=['PUT'])
@token_required
def update_layout(current_user, layout_id):
    """Memperbarui informasi metadata dari sebuah layout."""
    layout = Layout.query.get_or_404(layout_id)
    
    # Memastikan pengguna hanya bisa mengubah layout miliknya
    if layout.uploaded_by != current_user.id:
        return jsonify({"msg": "Akses ditolak. Anda bukan pemilik layout ini."}), 403
        
    data = request.get_json()
    if not data:
        return jsonify({"msg": "Request body tidak boleh kosong."}), 400
        
    # Memperbarui data jika ada di request
    layout.jenjang = data.get('jenjang', layout.jenjang)
    layout.mapel = data.get('mapel', layout.mapel)
    layout.tipe_dokumen = data.get('tipe_dokumen', layout.tipe_dokumen)
    
    db.session.commit()
    
    return jsonify({"msg": f"Layout dengan ID {layout.id} berhasil diperbarui."}), 200

# (DELETE) - Menghapus sebuah layout
@layout_bp.route('/<int:layout_id>', methods=['DELETE'])
@token_required
def delete_layout(current_user, layout_id):
    """Menghapus sebuah layout dari database dan file dari server."""
    layout = Layout.query.get_or_404(layout_id)
    
    # Memastikan pengguna hanya bisa menghapus layout miliknya
    if layout.uploaded_by != current_user.id:
        return jsonify({"msg": "Akses ditolak. Anda bukan pemilik layout ini."}), 403

    # Hapus file fisik dari server untuk menghemat ruang
    try:
        if os.path.exists(layout.file_path):
            os.remove(layout.file_path)
    except OSError as e:
        # Jika gagal menghapus file, kirim pesan error tapi tetap lanjutkan proses
        print(f"Error saat menghapus file {layout.file_path}: {e}")

    # Hapus data dari database
    db.session.delete(layout)
    db.session.commit()
    
    return jsonify({"msg": f"Layout dengan ID {layout.id} berhasil dihapus."}), 200